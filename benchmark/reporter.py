"""
Reporter — Reduction Experiment
================================
Generates reports and charts for the reduction experiment.

Outputs:
  - Console tables (rich): per-model reducer comparison
  - Charts (matplotlib):
            * <model>_hallucination_scores.png  - baseline vs reducers (all detectors)
            * <model>_per_detector.png          - one subplot per detector
            * <model>_score_reductions.png      - reduction vs baseline (positive is good)
            * <model>_latency.png               - mean latency per reducer
            * overall_hallucination_scores.png  - averaged across all models
            * overall_per_detector.png          - averaged across all models, per-detector
            * overall_score_reductions.png      - averaged across all models
            * overall_latency.png               - averaged across all models
    - HTML report: report.html (loads PNG charts from the same output folder)
    - DOCX report (optional): report.docx with tables + embedded PNG charts
  - JSON summary
"""
from __future__ import annotations

import json
from pathlib import Path
from typing import Optional

import pandas as pd
from loguru import logger
from rich.console import Console
from rich.table import Table
from rich import box
import numpy as np

console = Console()

# Score column display names
DETECTOR_LABELS = {
    "token_score":     "Token Similarity",
    "semantic_score":  "Semantic Similarity",
    "bert_score":      "BERT Stochastic",
    "llm_score":       "LLM Judge",
}

# Reducer display names
REDUCER_LABELS = {
    "baseline":              "Baseline (no reducer)",
    "rag":                   "RAG",
    "constrained_decoding":  "Constrained Decoding",
    "self_verification":     "Self-Verification",
}

# Colors for plots (red = baseline/bad, green = good reductions)
REDUCER_COLORS = {
    "baseline":              "#E74C3C",  # red
    "rag":                   "#2ECC71",  # green
    "constrained_decoding":  "#3498DB",  # blue
    "self_verification":     "#F39C12",  # orange
}


class Reporter:
    """Generates console tables, charts, and HTML reports for the reduction experiment."""

    def __init__(self, output_dir: str = "results"):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ══════════════════════════════════════════════════════════
    # Console output
    # ══════════════════════════════════════════════════════════

    def print_summary(self, summary_df: pd.DataFrame):
        """Print per-model reducer comparison table to the console."""
        if summary_df is None or len(summary_df) == 0:
            console.print("[yellow]No summary data to report.[/yellow]")
            return

        console.rule("[bold cyan]Reduction Experiment Results[/bold cyan]")

        score_cols = [c for c in DETECTOR_LABELS if c in summary_df.columns]

        for model_name, mdf in summary_df.groupby("model"):
            table = Table(
                title=f"Model: [bold green]{model_name}[/bold green]",
                box=box.ROUNDED,
                show_header=True,
                header_style="bold magenta",
                caption="[dim]Lower scores = less hallucination. Baseline first.[/dim]",
            )
            table.add_column("Reducer", style="cyan", width=24)
            table.add_column("Dataset", style="white", width=18)
            for col in score_cols:
                table.add_column(
                    DETECTOR_LABELS[col],
                    justify="right",
                    style="yellow",
                )
            table.add_column("N", justify="right", style="dim")

            # Order so baseline comes first
            reducer_order = ["baseline"] + [
                r for r in mdf["reducer"].unique() if r != "baseline"
            ]
            mdf_sorted = mdf.copy()
            mdf_sorted["_order"] = mdf_sorted["reducer"].map(
                {r: i for i, r in enumerate(reducer_order)}
            )
            mdf_sorted = mdf_sorted.sort_values(["_order", "dataset"])

            for _, row in mdf_sorted.iterrows():
                reducer_label = REDUCER_LABELS.get(row["reducer"], row["reducer"])
                cells = [reducer_label, str(row.get("dataset", "-"))]
                for col in score_cols:
                    val = row.get(col)
                    cells.append(f"{val:.3f}" if pd.notna(val) else "-")
                cells.append(str(int(row.get("n_samples", 0))))
                table.add_row(*cells)

            console.print(table)

    def print_reduction_table(self, reduction_df: pd.DataFrame):
        """Print the reduction-vs-baseline table."""
        if reduction_df is None or len(reduction_df) == 0:
            return

        console.rule("[bold yellow]Score Reduction vs Baseline[/bold yellow]")
        console.print("[dim]Positive = reducer lowered hallucination score. Higher = better.[/dim]\n")

        reduction_cols = [c for c in reduction_df.columns if c.endswith("_reduction")]

        for model_name, mdf in reduction_df.groupby("model"):
            table = Table(
                title=f"Model: [bold green]{model_name}[/bold green]",
                box=box.ROUNDED,
                header_style="bold magenta",
            )
            table.add_column("Reducer", style="cyan")
            for col in reduction_cols:
                label = DETECTOR_LABELS.get(
                    col.replace("_reduction", "_score"),
                    col.replace("_reduction", ""),
                )
                table.add_column(label, justify="right")

            for _, row in mdf.iterrows():
                reducer_label = REDUCER_LABELS.get(row["reducer"], row["reducer"])
                cells = [reducer_label]
                for col in reduction_cols:
                    val = row.get(col)
                    if pd.notna(val):
                        # Color: green if positive (good), red if negative (bad)
                        color = "green" if val > 0 else "red"
                        cells.append(f"[{color}]{val:+.3f}[/{color}]")
                    else:
                        cells.append("-")
                table.add_row(*cells)

            console.print(table)

    # ══════════════════════════════════════════════════════════
    # Charts
    # ══════════════════════════════════════════════════════════

    def save_charts(self, summary_df: pd.DataFrame, reduction_df: pd.DataFrame = None):
        """Save one set of charts per model, plus overall charts across all models."""
        import matplotlib
        matplotlib.use("Agg")

        import matplotlib.pyplot as plt

        # Brighter, easier-to-read defaults
        try:
            plt.style.use("seaborn-v0_8-whitegrid")
        except Exception:
            pass
        plt.rcParams.update({
            "figure.facecolor": "white",
            "axes.facecolor": "white",
            "savefig.facecolor": "white",
            "axes.titlesize": 14,
            "axes.labelsize": 12,
            "xtick.labelsize": 10,
            "ytick.labelsize": 10,
            "legend.fontsize": 10,
        })

        score_cols = [c for c in DETECTOR_LABELS if c in summary_df.columns]
        if not score_cols:
            logger.warning("No detector scores to plot.")
            return

        # Generate one set of charts per model
        for model_name, model_df in summary_df.groupby("model"):
            logger.info(f"  Generating charts for model: {model_name}")

            # Chart 1: Raw scores — baseline vs reducers across all detectors
            self._plot_before_after(model_df, score_cols, model_name)

            # Chart 2: Per detector subplots
            self._plot_per_detector_subplots(model_df, score_cols, model_name)

            # Chart 2b: Latency / performance (if available)
            if "mean_latency_s" in model_df.columns:
                self._plot_latency(model_df, model_name)

            # Chart 3: Score reductions vs baseline
            if reduction_df is not None and len(reduction_df) > 0:
                model_reduction = reduction_df[
                    reduction_df["model"] == model_name
                ].copy()
                if len(model_reduction) > 0:
                    self._plot_reductions(model_reduction, model_name)

        # Overall charts (averaged across all models/datasets)
        logger.info("  Generating overall charts (all models)")
        self._plot_before_after(summary_df, score_cols, "overall")
        self._plot_per_detector_subplots(summary_df, score_cols, "overall")
        if "mean_latency_s" in summary_df.columns:
            self._plot_latency(summary_df, "overall")
        if reduction_df is not None and len(reduction_df) > 0:
            self._plot_reductions(reduction_df, "overall")

        logger.info(f"✓ All charts saved to {self.output_dir}")

    def _plot_before_after(self, summary_df: pd.DataFrame, score_cols: list, model_name: str):
        """
        Grouped bar chart for one model:
        - Groups = detectors (Token, Semantic, BERT, LLM)
        - Bars within each group = reducers (Baseline, RAG, Constrained, Self-Verify)
        
        Shows raw hallucination scores before and after each reducer,
        scored by all detectors. Lower = less hallucination = better.
        """
        import matplotlib.pyplot as plt
        import numpy as np

        # Average scores per reducer for this model
        avg = summary_df.groupby("reducer")[score_cols].mean().round(4)

        # Always show baseline first
        reducer_order = ["baseline"] + [r for r in avg.index if r != "baseline"]
        avg = avg.loc[[r for r in reducer_order if r in avg.index]]

        fig, ax = plt.subplots(figsize=(14, 7))
        x = np.arange(len(score_cols))
        width = 0.8 / len(avg)

        for i, (reducer, row) in enumerate(avg.iterrows()):
            offset = (i - len(avg) / 2 + 0.5) * width
            bars = ax.bar(
                x + offset,
                row.values,
                width,
                label=REDUCER_LABELS.get(reducer, reducer),
                color=REDUCER_COLORS.get(reducer, "#999999"),
                edgecolor="white",
                linewidth=1.2,
            )
            for bar, val in zip(bars, row.values):
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + 0.01,
                    f"{val:.2f}",
                    ha="center", va="bottom",
                    fontsize=9, fontweight="bold",
                )

        ax.set_xticks(x)
        ax.set_xticklabels(
            [DETECTOR_LABELS[c] for c in score_cols],
            fontsize=11
        )
        ax.set_ylabel("Hallucination Score (lower = better)", fontsize=12)
        ax.set_title(
            f"Model: {model_name}\n"
            f"Hallucination Scores — Baseline vs Reducers (all detectors)",
            fontsize=14, fontweight="bold", pad=15,
        )
        ax.set_ylim(0, 1.15)
        ax.axhline(
            y=0.5, color="gray", linestyle="--",
            alpha=0.4, label="0.5 threshold"
        )
        ax.legend(loc="upper right", framealpha=0.9)
        ax.grid(axis="y", alpha=0.3, linestyle="--")

        # Safe filename — remove special chars from model name
        safe_name = model_name.replace(":", "_").replace("/", "_").replace(" ", "_")
        plt.tight_layout()
        path = self.output_dir / f"{safe_name}_hallucination_scores.png"
        plt.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close()
        logger.info(f"  Saved: {path}")

    def _plot_per_detector_subplots(self, summary_df: pd.DataFrame, score_cols: list, model_name: str):
        """
        One subplot per detector for a specific model.
        Each subplot shows scores per reducer (Baseline, RAG, Constrained, Self-Verify)
        for that detector specifically.
        """
        import matplotlib.pyplot as plt

        n = len(score_cols)
        cols = min(n, 2)
        rows = (n + cols - 1) // cols

        fig, axes = plt.subplots(rows, cols, figsize=(6 * cols, 5 * rows))
        if rows * cols == 1:
            axes = [axes]
        else:
            axes = axes.flatten()

        avg = summary_df.groupby("reducer")[score_cols].mean().round(4)
        reducer_order = ["baseline"] + [r for r in avg.index if r != "baseline"]
        avg = avg.loc[[r for r in reducer_order if r in avg.index]]

        for i, col in enumerate(score_cols):
            ax = axes[i]
            reducers = avg.index.tolist()
            values   = avg[col].values
            colors   = [REDUCER_COLORS.get(r, "#999999") for r in reducers]
            labels   = [REDUCER_LABELS.get(r, r) for r in reducers]

            bars = ax.bar(labels, values, color=colors,
                        edgecolor="white", linewidth=1.2, width=0.6)
            for bar, val in zip(bars, values):
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    bar.get_height() + 0.01,
                    f"{val:.3f}",
                    ha="center", va="bottom",
                    fontsize=10, fontweight="bold",
                )

            # Each subplot title = detector name
            ax.set_title(
                f"{DETECTOR_LABELS[col]}\nModel: {model_name}",
                fontsize=11, fontweight="bold"
            )
            ax.set_ylabel("Hallucination Score (lower = better)")
            ax.set_ylim(0, 1.15)
            ax.axhline(y=0.5, color="gray", linestyle="--", alpha=0.4)
            ax.grid(axis="y", alpha=0.3, linestyle="--")
            ax.tick_params(axis="x", rotation=20)

        # Hide unused subplots if detectors < grid size
        for j in range(len(score_cols), len(axes)):
            axes[j].axis("off")

        fig.suptitle(
            f"Model: {model_name}\n"
            f"Per-Detector Hallucination Scores — Baseline vs Reducers",
            fontsize=14, fontweight="bold", y=1.02,
        )
        plt.tight_layout()

        # Save with model name in filename
        safe_name = model_name.replace(":", "_").replace("/", "_").replace(" ", "_")
        path = self.output_dir / f"{safe_name}_per_detector.png"
        plt.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close()
        logger.info(f"  Saved: {path}")

    def _plot_reductions(self, reduction_df: pd.DataFrame, model_name: str):
        """
        Grouped bar chart for one model showing score REDUCTION vs baseline.
        
        - Groups = detectors (Token, Semantic, BERT, LLM)
        - Bars within each group = reducers (RAG, Constrained, Self-Verify)
        
        Positive value = reducer lowered hallucination score (good).
        Negative value = reducer made things worse (bad).
        """
        import matplotlib.pyplot as plt
        import numpy as np

        reduction_cols = [c for c in reduction_df.columns if c.endswith("_reduction")]
        if not reduction_cols:
            return

        avg = reduction_df.groupby("reducer")[reduction_cols].mean().round(4)

        fig, ax = plt.subplots(figsize=(12, 6))
        x = np.arange(len(reduction_cols))
        width = 0.8 / len(avg)

        for i, (reducer, row) in enumerate(avg.iterrows()):
            offset = (i - len(avg) / 2 + 0.5) * width
            bars = ax.bar(
                x + offset,
                row.values,
                width,
                label=REDUCER_LABELS.get(reducer, reducer),
                color=REDUCER_COLORS.get(reducer, "#999999"),
                edgecolor="white",
                linewidth=1.2,
            )
            for bar, val in zip(bars, row.values):
                y_pos = bar.get_height()
                va = "bottom" if y_pos >= 0 else "top"
                offset_y = 0.005 if y_pos >= 0 else -0.005
                ax.text(
                    bar.get_x() + bar.get_width() / 2,
                    y_pos + offset_y,
                    f"{val:+.2f}",
                    ha="center", va=va,
                    fontsize=9,
                )

        ax.axhline(0, color="black", linewidth=1.2)
        detector_labels = [
            DETECTOR_LABELS.get(c.replace("_reduction", "_score"), c)
            for c in reduction_cols
        ]
        ax.set_xticks(x)
        ax.set_xticklabels(detector_labels, fontsize=11)
        ax.set_ylabel(
            "Score Reduction vs Baseline\n(positive = less hallucination)",
            fontsize=12
        )
        ax.set_title(
            f"Model: {model_name}\n"
            f"How Much Each Reducer Lowered Hallucination Scores",
            fontsize=14, fontweight="bold", pad=15,
        )
        ax.legend(loc="upper right", framealpha=0.9)
        ax.grid(axis="y", alpha=0.3, linestyle="--")

        safe_name = model_name.replace(":", "_").replace("/", "_").replace(" ", "_")
        plt.tight_layout()
        path = self.output_dir / f"{safe_name}_score_reductions.png"
        plt.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close()
        logger.info(f"  Saved: {path}")

    def _plot_latency(self, summary_df: pd.DataFrame, model_name: str):
        """Bar chart of mean latency per reducer (averaged across datasets)."""
        import matplotlib.pyplot as plt
        if "mean_latency_s" not in summary_df.columns:
            return

        avg = (
            summary_df.groupby("reducer")["mean_latency_s"]
                      .mean()
                      .round(4)
        )
        if len(avg) == 0:
            return

        reducer_order = ["baseline"] + [r for r in avg.index if r != "baseline"]
        avg = avg.loc[[r for r in reducer_order if r in avg.index]]

        fig, ax = plt.subplots(figsize=(11, 5))
        labels = [REDUCER_LABELS.get(r, r) for r in avg.index]
        colors = [REDUCER_COLORS.get(r, "#999999") for r in avg.index]
        bars = ax.bar(labels, avg.values, color=colors, edgecolor="white", linewidth=1.2)

        for bar, val in zip(bars, avg.values):
            ax.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height() + (max(avg.values) * 0.02 if max(avg.values) else 0.01),
                f"{val:.2f}s",
                ha="center", va="bottom",
                fontsize=10, fontweight="bold",
            )

        title_name = "Overall (all models)" if model_name == "overall" else f"Model: {model_name}"
        ax.set_title(
            f"{title_name}\nMean Latency per Answer — Baseline vs Reducers",
            fontsize=14, fontweight="bold", pad=15,
        )
        ax.set_ylabel("Mean latency (seconds)")
        ax.grid(axis="y", alpha=0.3, linestyle="--")
        ax.tick_params(axis="x", rotation=15)

        safe_name = model_name.replace(":", "_").replace("/", "_").replace(" ", "_")
        plt.tight_layout()
        path = self.output_dir / f"{safe_name}_latency.png"
        plt.savefig(path, dpi=200, bbox_inches="tight", facecolor="white")
        plt.close()
        logger.info(f"  Saved: {path}")
    # ══════════════════════════════════════════════════════════
    # Output files
    # ══════════════════════════════════════════════════════════

    def save_json_summary(self, summary_df: pd.DataFrame,
                          reduction_df: Optional[pd.DataFrame] = None):
        """Save JSON summary of the experiment."""
        out = {
            "summary": summary_df.to_dict(orient="records") if summary_df is not None else [],
        }
        if reduction_df is not None and len(reduction_df) > 0:
            out["reductions"] = reduction_df.to_dict(orient="records")

        path = self.output_dir / "summary.json"
        with open(path, "w") as f:
            json.dump(out, f, indent=2, default=str)
        logger.info(f"  Saved: {path}")

    def save_html_report(self, summary_df: pd.DataFrame,
                     reduction_df: Optional[pd.DataFrame] = None,
                     results_df: Optional[pd.DataFrame] = None):
        """Save an HTML report with per-model charts and tables.

        Note: the HTML references PNG files in the same folder (it does not embed images).
        """
        html_parts = [
            "<!DOCTYPE html><html><head>",
            "<meta charset='utf-8'>",
            "<title>Hallucination Reduction Experiment Report</title>",
            "<style>",
            "body { font-family: -apple-system, Arial, sans-serif; max-width: 1200px; "
            "margin: 2em auto; padding: 0 1em; color: #333; }",
            "h1 { border-bottom: 3px solid #3498DB; padding-bottom: 10px; }",
            "h2 { color: #2C3E50; margin-top: 2em; }",
            "h3 { color: #7F8C8D; }",
            "table { border-collapse: collapse; margin: 1em 0; width: 100%; }",
            "th { background: #3498DB; color: white; padding: 8px 12px; text-align: left; }",
            "td { padding: 6px 12px; border-bottom: 1px solid #ddd; }",
            "tr:nth-child(even) { background: #f9f9f9; }",
            "img { max-width: 100%; margin: 1em 0; border: 1px solid #ddd; "
            "border-radius: 4px; }",
            ".caption { color: #666; font-size: 0.9em; margin: 0.5em 0; }",
            ".model-section { border: 1px solid #ddd; border-radius: 8px; "
            "padding: 1em; margin: 2em 0; }",
            "</style></head><body>",
            "<h1>Hallucination Reduction Experiment Report</h1>",
            "<p class='caption'>Testing how RAG, Constrained Decoding, and "
            "Self-Verification reduce LLM hallucination scores, "
            "measured by 4 detectors (Token, Semantic, BERT, LLM Judge).</p>",
        ]

        # Overall charts (if present)
        overall_charts = [
            ("overall_hallucination_scores.png", "Overall hallucination scores"),
            ("overall_score_reductions.png", "Overall score reductions"),
            ("overall_latency.png", "Overall latency"),
            ("overall_per_detector.png", "Overall per-detector scores"),
        ]
        if any((self.output_dir / p).exists() for p, _ in overall_charts):
            html_parts.append("<h2>Overall Results (All Models)</h2>")
            html_parts.append(
                "<p class='caption'>Averaged across all models/datasets in this run folder.</p>"
            )
            for p, alt in overall_charts:
                if (self.output_dir / p).exists():
                    html_parts.append(f"<img src='{p}' alt='{alt}'>")

        # Per-model sections
        if summary_df is not None and len(summary_df) > 0:
            html_parts.append("<h2>Per-Model Results</h2>")

            for model_name in summary_df["model"].unique():
                safe_name = model_name.replace(":", "_").replace(
                    "/", "_").replace(" ", "_")
                html_parts.append(f"<div class='model-section'>")
                html_parts.append(f"<h3>Model: {model_name}</h3>")

                # Scores chart
                scores_chart = f"{safe_name}_hallucination_scores.png"
                if (self.output_dir / scores_chart).exists():
                    html_parts.append(
                        f"<img src='{scores_chart}' "
                        f"alt='Hallucination scores for {model_name}'>"
                    )

                # Reductions chart
                reductions_chart = f"{safe_name}_score_reductions.png"
                if (self.output_dir / reductions_chart).exists():
                    html_parts.append(
                        f"<img src='{reductions_chart}' "
                        f"alt='Score reductions for {model_name}'>"
                    )

                # Per-detector chart
                detector_chart = f"{safe_name}_per_detector.png"
                if (self.output_dir / detector_chart).exists():
                    html_parts.append(
                        f"<img src='{detector_chart}' "
                        f"alt='Per-detector scores for {model_name}'>"
                    )

                # Latency chart
                latency_chart = f"{safe_name}_latency.png"
                if (self.output_dir / latency_chart).exists():
                    html_parts.append(
                        f"<img src='{latency_chart}' "
                        f"alt='Latency per reducer for {model_name}'>"
                    )

                # Per-model summary table
                model_summary = summary_df[summary_df["model"] == model_name]
                score_cols = [c for c in DETECTOR_LABELS if c in model_summary.columns]
                display_cols = ["reducer", "dataset"] + score_cols + ["n_samples"]
                display_cols = [c for c in display_cols if c in model_summary.columns]
                html_parts.append(
                    model_summary[display_cols].to_html(index=False, float_format="%.3f")
                )
                html_parts.append("</div>")

        # Overall reduction table across all models
        if reduction_df is not None and len(reduction_df) > 0:
            html_parts.append("<h2>Score Reductions vs Baseline (All Models)</h2>")
            html_parts.append(
                "<p class='caption'>Positive = reducer lowered hallucination score. "
                "Higher = better.</p>"
            )
            html_parts.append(
                reduction_df.to_html(index=False, float_format="%.3f")
            )

        html_parts.append("</body></html>")

        path = self.output_dir / "report.html"
        with open(path, "w") as f:
            f.write("\n".join(html_parts))
        logger.info(f"  Saved: {path}")

    def save_docx_report(
        self,
        summary_df: pd.DataFrame,
        reduction_df: Optional[pd.DataFrame] = None,
        results_df: Optional[pd.DataFrame] = None,
        filename: str = "report.docx",
    ):
        """Save a Word (.docx) report with tables and embedded PNG charts.

        This is intended for sharing (one file containing tables + figures).
        Requires the optional dependency: python-docx.
        """
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception as e:
            logger.error(
                "python-docx is required for DOCX export. "
                "Install it with: pip install python-docx\n"
                f"Import error: {e}"
            )
            return

        def _safe(name: str) -> str:
            return name.replace(":", "_").replace("/", "_").replace(" ", "_")

        def _add_picture(doc: Document, rel_path: str, caption: Optional[str] = None):
            p = self.output_dir / rel_path
            if not p.exists():
                return
            if caption:
                doc.add_paragraph(caption)
            # 6.5" fits comfortably on US Letter with default margins
            doc.add_picture(str(p), width=Inches(6.5))

        def _add_df_table(doc: Document, df: pd.DataFrame):
            if df is None or len(df) == 0:
                doc.add_paragraph("(no data)")
                return

            display_df = df.copy()
            for col in display_df.columns:
                if np.issubdtype(display_df[col].dtype, np.number):
                    display_df[col] = display_df[col].round(4)
            display_df = display_df.replace({np.nan: ""})

            table = doc.add_table(rows=1, cols=len(display_df.columns))
            table.style = "Table Grid"

            # Header row
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(display_df.columns):
                hdr_cells[i].text = str(col)
                if hdr_cells[i].paragraphs and hdr_cells[i].paragraphs[0].runs:
                    hdr_cells[i].paragraphs[0].runs[0].bold = True

            # Data rows
            for _, row in display_df.iterrows():
                cells = table.add_row().cells
                for i, col in enumerate(display_df.columns):
                    val = row[col]
                    cells[i].text = "" if val is None else str(val)

        doc = Document()
        doc.add_heading("Hallucination Reduction Experiment Report", level=0)
        doc.add_paragraph(
            "This report includes tables and figures generated by the benchmark. "
            "Lower detector scores mean less hallucination."
        )

        # Overall charts
        overall_charts = [
            ("overall_hallucination_scores.png", "Overall hallucination scores"),
            ("overall_per_detector.png", "Overall per-detector scores"),
            ("overall_score_reductions.png", "Overall score reductions vs baseline"),
            ("overall_latency.png", "Overall latency"),
        ]
        if any((self.output_dir / p).exists() for p, _ in overall_charts):
            doc.add_heading("Overall Results (All Models)", level=1)
            for rel, caption in overall_charts:
                _add_picture(doc, rel, caption)

        # Per-model sections
        if summary_df is not None and len(summary_df) > 0 and "model" in summary_df.columns:
            doc.add_heading("Per-Model Results", level=1)
            for model_name in summary_df["model"].unique():
                safe = _safe(str(model_name))
                doc.add_heading(f"Model: {model_name}", level=2)
                _add_picture(doc, f"{safe}_hallucination_scores.png")
                _add_picture(doc, f"{safe}_per_detector.png")
                _add_picture(doc, f"{safe}_score_reductions.png")
                _add_picture(doc, f"{safe}_latency.png")

        # Tables
        doc.add_heading("Summary Table (Mean Scores)", level=1)
        _add_df_table(doc, summary_df)

        if reduction_df is not None and len(reduction_df) > 0:
            doc.add_heading("Reductions vs Baseline", level=1)
            _add_df_table(doc, reduction_df)

        out_path = self.output_dir / filename
        doc.save(str(out_path))
        logger.info(f"  Saved: {out_path}")